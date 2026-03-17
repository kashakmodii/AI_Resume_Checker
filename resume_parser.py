"""
Resume Parser Module
Handles extraction of text from PDF and DOCX resume files.
"""

import PyPDF2
from docx import Document
import streamlit as st
import os
from pathlib import Path


def extract_text_from_pdf(file_input):
    """
    Extract text from a PDF file.
    
    Args:
        file_input: Path to the PDF file or file-like object (UploadedFile)
        
    Returns:
        Extracted text as string
        
    Raises:
        Exception: If PDF parsing fails
    """
    try:
        text = ""
        # Check if it's a file path or file-like object
        if isinstance(file_input, str):
            with open(file_input, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
        else:
            # Handle Streamlit UploadedFile or file-like objects
            reader = PyPDF2.PdfReader(file_input)
            for page in reader.pages:
                text += page.extract_text()
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_docx(file_input):
    """
    Extract text from a DOCX (Word) file.
    
    Args:
        file_input: Path to the DOCX file or file-like object (UploadedFile)
        
    Returns:
        Extracted text as string
        
    Raises:
        Exception: If DOCX parsing fails
    """
    try:
        # Check if it's a file path or file-like object
        if isinstance(file_input, str):
            doc = Document(file_input)
        else:
            # Handle Streamlit UploadedFile or file-like objects
            doc = Document(file_input)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


def extract_text_from_upload(uploaded_file):
    """
    Extract text from an uploaded resume file (PDF or DOCX).
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If file format is not supported
        Exception: If text extraction fails
    """
    try:
        if uploaded_file.type == "application/pdf":
            with st.spinner("🔄 Extracting text from PDF..."):
                return extract_text_from_pdf(uploaded_file)
        
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with st.spinner("🔄 Extracting text from Word document..."):
                return extract_text_from_docx(uploaded_file)
        
        else:
            raise ValueError(f"Unsupported file type: {uploaded_file.type}. Please upload PDF or DOCX.")
    
    except Exception as e:
        raise Exception(f"Failed to extract resume text: {str(e)}")


def load_job_descriptions(file_path):
    """
    Load job descriptions from a text file with format: Job Title: description
    Robust path handling with fallback to embedded data for Streamlit Cloud
    
    Args:
        file_path: Path to the job descriptions file
        
    Returns:
        Dictionary with job titles as keys and descriptions as values
    """
    job_descriptions = {}

    try:
        # Try multiple path resolution strategies
        possible_paths = [
            file_path,  # Direct path as given
            Path(file_path),  # Using pathlib
            Path(__file__).parent / file_path,  # Relative to this module
        ]
        
        # Add current working directory variants
        cwd = Path.cwd()
        possible_paths.extend([
            cwd / file_path,
            cwd / "job_description.txt",  # Default filename
        ])
        
        # Try to find the file
        actual_path = None
        for p in possible_paths:
            if isinstance(p, str):
                p = Path(p)
            if p.exists():
                actual_path = p
                break
        
        if actual_path is not None:
            with open(actual_path, 'r', encoding='utf-8') as file:
                content = file.read()
                lines = content.split('\n')
                
                current_title = ""
                current_desc = []
                
                for line in lines:
                    line_stripped = line.strip()
                    
                    # Check if this line is a job title (ends with colon)
                    if line_stripped and ':' in line_stripped and not line_stripped.startswith('#'):
                        # Save previous job if exists
                        if current_title and current_desc:
                            job_descriptions[current_title] = ' '.join(current_desc).strip()
                        
                        # Parse new job title
                        parts = line_stripped.split(':', 1)
                        current_title = parts[0].strip()
                        current_desc = []
                        
                        # If there's text after the colon on the same line, add it
                        if len(parts) > 1 and parts[1].strip():
                            current_desc.append(parts[1].strip())
                    
                    elif line_stripped and current_title:
                        # Add to current description
                        current_desc.append(line_stripped)
                
                # Don't forget the last job
                if current_title and current_desc:
                    job_descriptions[current_title] = ' '.join(current_desc).strip()
        
        # If file not found or empty, use fallback embedded data
        if not job_descriptions:
            job_descriptions = get_fallback_job_descriptions()
        
        return job_descriptions
    
    except Exception as e:
        # On any error, return fallback data instead of failing
        return get_fallback_job_descriptions()


def get_fallback_job_descriptions():
    """
    Fallback job descriptions embedded in the app.
    Used when job_description.txt is not available or cannot be parsed.
    Ensures templates always load on Streamlit Cloud.
    """
    return {
        "Data Scientist": "Responsible for collecting, analyzing, and interpreting large datasets to generate actionable insights. Uses statistical techniques, machine learning models, and data visualization tools to support decision-making and business strategies.",
        "Software Engineer": "Designs, develops, tests, and maintains software applications. Works with programming languages, frameworks, and development tools to build scalable and efficient systems.",
        "Web Developer": "Builds and maintains websites and web applications. Handles frontend and backend development, ensuring responsive design, performance optimization, and cross-browser compatibility.",
        "Machine Learning Engineer": "Develops and deploys machine learning models into production. Works on data pipelines, model optimization, and scalable AI solutions.",
        "Android Developer": "Designs and develops mobile applications for Android platforms. Ensures performance, usability, and integration with backend services.",
        "Data Analyst": "Collects, processes, and analyzes data to identify trends and patterns. Creates reports and dashboards to assist business decisions.",
        "UI/UX Designer": "Designs user-friendly interfaces and enhances user experience. Conducts research, wireframing, prototyping, and usability testing.",
        "Network Engineer": "Designs, implements, and manages computer networks. Ensures network security, performance, and reliability.",
        "DevOps Engineer": "Bridges development and operations by automating workflows, managing CI/CD pipelines, and ensuring system scalability and reliability.",
        "Cybersecurity Analyst": "Protects systems and networks from cyber threats. Monitors security incidents, conducts vulnerability assessments, and implements security measures.",
        "Project Manager": "Plans, executes, and manages projects within scope, budget, and timeline. Coordinates teams and ensures successful project delivery.",
        "Graphic Designer": "Creates visual content for branding, marketing, and digital platforms. Uses design tools to produce engaging graphics.",
        "Content Writer": "Develops written content for websites, blogs, and marketing materials. Focuses on clarity, engagement, and SEO optimization.",
        "Marketing Manager": "Plans and executes marketing strategies to promote products or services. Analyzes market trends and manages campaigns.",
        "Accountant": "Manages financial records, prepares reports, and ensures compliance with regulations. Handles budgeting, auditing, and taxation.",
        "Human Resources Specialist": "Manages recruitment, employee relations, and organizational policies. Supports workforce development and performance management.",
        "Customer Support Representative": "Assists customers by resolving queries and issues. Ensures customer satisfaction through effective communication.",
        "Mechanical Engineer": "Designs and develops mechanical systems and components. Works with CAD software, prototyping, and testing to create innovative solutions.",
        "Electrical Engineer": "Designs and develops electrical systems and equipment. Manages power distribution, circuit design, and electrical installations.",
        "Civil Engineer": "Plans and designs infrastructure projects like buildings, bridges, and roads. Ensures safety, sustainability, and compliance with regulations.",
        "AI Engineer": "Develops artificial intelligence solutions using deep learning, NLP, and computer vision. Deploys AI models for real-world applications.",
        "Database Administrator": "Manages databases, ensures data security and integrity. Performs backups, optimization, and maintenance of database systems.",
        "Cloud Engineer": "Designs and manages cloud infrastructure on platforms like AWS, Azure, or GCP. Ensures scalability, security, and cost-efficiency.",
        "Frontend Developer": "Builds user interfaces for web applications using HTML, CSS, JavaScript, and modern frameworks. Ensures responsive and interactive designs.",
        "Backend Developer": "Develops server-side logic, APIs, and databases. Works with databases, servers, and backend frameworks to support applications.",
        "Full Stack Developer": "Handles both frontend and backend development. Works across the entire application stack to deliver complete solutions.",
        "QA Tester": "Tests software applications to identify bugs and ensure quality. Creates test cases, performs manual and automated testing.",
    }
